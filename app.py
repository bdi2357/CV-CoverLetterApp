import streamlit as st
import os
from docx import Document
from PyPDF2 import PdfReader
import openai
import logging
##
import os
print("Available files: ", os.listdir())
print("Available in CoverLetterGen: ", os.listdir('CoverLetterGen'))

from CoverLetterGen.ai_interaction import OpenAIModel, CoverLetterGenerator
from CoverLetterGen.basic_iterative import BasicIterativeAgent

if "OPENAI_API_KEY" in st.secrets:
    st.write("API key loaded successfully.")  # This will be shown in the app
    print(f"OPENAI_API_KEY: {st.secrets['OPENAI_API_KEY']}")  # This will be printed in the logs
else:
    st.write("OPENAI_API_KEY not found in secrets.")  # If the key is not found
    print("OPENAI_API_KEY not found.")
# Verify the version of the OpenAI library
st.write(f"OpenAI library version: {openai.__version__}")  # Shows in the app
print(f"OpenAI library version: {openai.__version__}")  # Shows in the logs
logging.info(f"OpenAI library version: {openai.__version__}")
# Initialize the BasicIterativeAgent
ai_model = OpenAIModel(api_key=st.secrets["OPENAI_API_KEY"], model_name='gpt-4o')
cover_letter_gen = CoverLetterGenerator(ai_model)
agent = BasicIterativeAgent(cover_letter_gen)


def extract_text_from_pdf(pdf_file):
    """Extracts text from a PDF file."""
    reader = PdfReader(pdf_file)
    text = ''
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_file):
    """Extracts text from a DOCX file."""
    doc = Document(docx_file)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return text

def save_cover_letter_to_docx(cover_letter_text, file_name):
    """
    Saves the generated cover letter as a DOCX file.

    Parameters:
        cover_letter_text (str): The generated cover letter text.
        file_name (str): The name of the DOCX file to save.
    """
    doc = Document()
    doc.add_paragraph(cover_letter_text)
    os.makedirs("Output", exist_ok=True)
    file_path = os.path.join("Output", file_name)
    doc.save(file_path)
    return file_path

# Streamlit UI components
st.title('📄 AI Cover Letter Generator')

st.write("Upload your CV and paste the job description to generate a professional cover letter.")
st.write("")

# File uploader for CV (PDF or DOCX)
uploaded_cv = st.file_uploader('📁 Upload your CV (PDF or DOCX)', type=['pdf', 'docx'])

# Text area for job description
job_description = st.text_area('✍️ Paste the job description here:', height=200)

# Button to generate cover letter
if st.button('Generate Cover Letter'):
    if uploaded_cv is None or job_description.strip() == '':
        st.warning('⚠️ Please upload your CV and enter a job description.')
    else:
        st.info("✍️ Generating your cover letter... Please wait.")

        # Detect file type and extract text accordingly
        file_type = uploaded_cv.type
        if file_type == "application/pdf":
            cv_text = extract_text_from_pdf(uploaded_cv)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            cv_text = extract_text_from_docx(uploaded_cv)

        # Generate initial cover letter using BasicIterativeAgent
        cover_letter = agent.generate_cover_letter(cv_text, job_description)

        # Optionally improve the generated cover letter
        improved_cover_letter, final_critique = agent.improve_cover_letter(cv_text, cover_letter, job_description)

        # Save the improved cover letter as a DOCX file
        docx_file_path = save_cover_letter_to_docx(improved_cover_letter, "cover_letter.docx")
        critique_cover_file_path = save_cover_letter_to_docx(final_critique, "cover_letter_critique.docx")

        # Display the cover letter to the user
        st.success('✅ Cover Letter generated successfully!')
        st.download_button(
            label='📥 Download Cover Letter',
            data=open(docx_file_path, 'rb').read(),
            file_name='cover_letter.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        st.download_button(
            label='📥 Download Cover Letter Critique',
            data=open(critique_cover_file_path, 'rb').read(),
            file_name='cover_letter_critique.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
